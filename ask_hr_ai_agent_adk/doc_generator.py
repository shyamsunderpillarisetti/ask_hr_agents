import os
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
from io import BytesIO

from docx import Document
from jinja2 import Environment, StrictUndefined
try:
    from docxtpl import DocxTemplate  # Optional, used for templated docs
    _HAS_DOXCTPL = True
except Exception:
    DocxTemplate = None  # type: ignore
    _HAS_DOXCTPL = False


BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

# In-memory document cache: maps doc_key -> {"bytes": BytesIO, "filename": str}
_document_cache: Dict[str, Dict[str, Any]] = {}


def _sanitize_filename(name: str, preserve_spaces: bool = False) -> str:
    """Sanitize filename by removing only invalid filesystem characters.
    
    Args:
        name: Original filename
        preserve_spaces: If True, keep spaces; otherwise convert to underscores
    
    Returns:
        Sanitized filename
    """
    # Remove only invalid Windows filename characters: < > : " / \ | ? *
    invalid_chars = '<>:"/\\|?*'
    safe = "".join(c for c in name if c not in invalid_chars)
    safe = safe.strip()
    if not preserve_spaces:
        safe = safe.replace(" ", "_")
    return safe or f"document_{int(datetime.now().timestamp())}"


def generate_docx(title: str, content: str, filename: Optional[str] = None) -> Dict[str, str]:
    """Generate a simple DOCX document and return file metadata.

    Args:
        title: Document title (adds a heading)
        content: Main body text content
        filename: Optional desired filename (without path). If not provided, auto-generates.

    Returns:
        Dict with keys: success, filename, download_key
    """
    title = (title or "Document").strip()
    content = (content or "").strip()

    base_name = _sanitize_filename(filename or f"{title.lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    if not base_name.lower().endswith(".docx"):
        base_name += ".docx"

    doc = Document()
    doc.add_heading(title, level=1)
    if content:
        for line in content.splitlines():
            doc.add_paragraph(line)
    
    # Save to BytesIO instead of disk
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Store in cache with unique key
    doc_key = f"{datetime.now().timestamp()}_{base_name}"
    _document_cache[doc_key] = {"bytes": output, "filename": base_name}

    return {
        "success": True,
        "filename": base_name,
        "download_key": doc_key,
    }


def list_templates() -> List[str]:
    """List available .docx templates in the templates folder."""
    return sorted([p.name for p in TEMPLATES_DIR.glob("*.docx")])


def get_template_placeholders(template_name: str) -> List[str]:
    """Return a sorted list of placeholder variables used in the template.

    Tries docxtpl first; if parsing fails, falls back to regex scanning.
    """
    tpl_path = TEMPLATES_DIR / template_name
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {template_name}")

    if _HAS_DOXCTPL:
        try:
            tpl = DocxTemplate(str(tpl_path))
            vars_set = tpl.get_undeclared_template_variables()
            return sorted(list(vars_set))
        except Exception:
            pass

    doc = Document(str(tpl_path))
    return _regex_scan_placeholders(doc)


def generate_docx_from_template(template_name: str, context: Dict[str, Any], filename: Optional[str] = None) -> Dict[str, Any]:
    """Render a DOCX from a .docx template and store in memory.

    Args:
        template_name: Filename of the template inside templates/
        context: Dict of placeholders to values
        filename: Optional output filename

    Returns:
        Dict with keys: success, filename, download_key (for Flask route)
    """
    if not _HAS_DOXCTPL:
        raise RuntimeError("docxtpl is not installed. Please install 'docxtpl'.")

    tpl_path = TEMPLATES_DIR / template_name
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {template_name}")

    # Use preserve_spaces=True for user-friendly filenames like "Employment Verification Letter - John Doe.docx"
    preserve_spaces = filename is not None and " - " in filename
    base_name = _sanitize_filename(filename or f"render_{Path(template_name).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", preserve_spaces=preserve_spaces)
    if not base_name.lower().endswith(".docx"):
        base_name += ".docx"

    env = Environment(undefined=StrictUndefined, autoescape=False)
    tpl = DocxTemplate(str(tpl_path))
    tpl.render(context or {}, jinja_env=env)
    
    # Save to BytesIO instead of disk
    output = BytesIO()
    tpl.save(output)
    output.seek(0)
    
    # Store in cache with unique key
    doc_key = f"{datetime.now().timestamp()}_{base_name}"
    _document_cache[doc_key] = {"bytes": output, "filename": base_name}

    return {
        "success": True,
        "filename": base_name,
        "download_key": doc_key,
    }


def get_document_from_cache(doc_key: str) -> Optional[BytesIO]:
    """Retrieve a document from memory cache by key."""
    cache_entry = _document_cache.get(doc_key)
    if cache_entry:
        return cache_entry.get("bytes")
    return None


def get_document_filename_from_cache(doc_key: str) -> Optional[str]:
    """Retrieve the filename of a cached document by key."""
    cache_entry = _document_cache.get(doc_key)
    if cache_entry:
        return cache_entry.get("filename")
    return None


def clear_document_cache(doc_key: str):
    """Remove a document from memory cache."""
    _document_cache.pop(doc_key, None)
